from docx import Document
import os
import findWordFile
import sys

# Giải mã văn bản bằng khóa riêng
def decrypt(private_key, ciphertext):
    d, n = private_key
    decrypted_message = ''.join([chr((char ** d) % n) for char in ciphertext])
    return decrypted_message

# Đọc nội dung từ tài liệu Word
def read_docx(filename):
    doc = Document(filename)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Ghi nội dung vào tài liệu Word
def write_docx(filename, text):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)

# Giải mã tài liệu Word
def decrypt_docx(input_filename, output_filename, private_key):
    encrypted_text = read_docx(input_filename)
    encrypted_message = list(map(int, encrypted_text.split()))
    decrypted_message = decrypt(private_key, encrypted_message)
    write_docx(output_filename, decrypted_message)
    
findWordFile.exec() ## Tìm đường dẫn đến các file word

def resource_path(relative_path):
    # Trả về đường dẫn tuyệt đối tới các file tài nguyên, 
    # hỗ trợ cả khi chạy code và khi đóng gói bằng PyInstaller
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

###ENCRYPT###
path =os.path.dirname(os.path.abspath(__file__)) #đường dẫn đến thư mục

#IMPORT PRIVATE KEY 
# Mở file và đọc nội dung
with open(path+'/private_key.txt', 'r') as file:
    content = file.read().strip()  # Đọc và loại bỏ khoảng trắng hai bên
# Chuyển đổi chuỗi thành tuple
# Tách các số bằng khoảng trắng và chuyển từng phần tử thành số nguyên
private_key = tuple(map(int, content.split()))

###ENCRYPT###
# Đường dẫn tới tệp của bạn
file_path = path+"/word_files_list.txt"

# Khởi tạo danh sách để lưu các dòng
lines = []
# Mở tệp và đọc từng dòng
with open(file_path, "r") as f:
    # Đọc tất cả các dòng và loại bỏ ký tự xuống dòng ở cuối mỗi dòng
    lines = [line.strip() for line in f]
# In danh sách các đường dẫn
print(lines)

for line in lines:
    decrypt_docx(line,line,private_key)
print("Success")




